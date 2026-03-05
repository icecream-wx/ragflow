"""
文件解析：支持 txt、docx 等文本类型，提取纯文本
"""
import os
from typing import List, Tuple
from app.core.logger import app_logger


# 允许的扩展名
ALLOWED_EXTENSIONS = {".txt", ".docx", ".doc"}


def parse_file(content: bytes, filename: str) -> str:
    """
    根据文件类型解析内容，返回纯文本。
    content: 文件二进制内容
    filename: 原始文件名（用于判断类型）
    """
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型: {ext}，仅支持 {', '.join(ALLOWED_EXTENSIONS)}")

    if ext == ".txt":
        return _parse_txt(content)
    if ext in (".docx", ".doc"):
        return _parse_docx(content, ext)
    return ""


def _parse_txt(content: bytes) -> str:
    """解析 txt：尝试 utf-8，再尝试 gbk。"""
    for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return content.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace").strip()


def _parse_docx(content: bytes, ext: str) -> str:
    """解析 docx（.doc 也尝试用 python-docx 读，部分可成功）。"""
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(content))
        parts = []
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n\n".join(parts) if parts else ""
    except Exception as e:
        app_logger.warning(f"解析 docx 失败: {e}")
        if ext == ".doc":
            raise ValueError("仅支持 .docx 格式的 Word 文档，请另存为 .docx 后重试")
        raise ValueError(f"解析 Word 文档失败: {e}") from e


def parse_uploaded_files(files: List[Tuple[bytes, str]]) -> List[Tuple[str, str]]:
    """
    解析多个上传文件，返回 [(纯文本, 文件名), ...]
    """
    results = []
    for content, filename in files:
        try:
            text = parse_file(content, filename)
            if text:
                results.append((text, filename))
            else:
                app_logger.warning(f"文件 {filename} 解析后为空，已跳过")
        except Exception as e:
            app_logger.error(f"解析 {filename} 失败: {e}")
            raise
    return results
